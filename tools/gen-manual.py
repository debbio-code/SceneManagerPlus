# -*- coding: utf-8 -*-
"""Genera il manuale utente di Scene Manager+ in formato Word."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x2E, 0x75, 0xB6)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
HEAD_FILL = "2E75B6"
ALT_FILL = "EAF1F8"

doc = Document()

# --- Stili base ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
    st = doc.styles[lvl]
    st.font.name = 'Calibri'
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = ACCENT if lvl != 'Heading 3' else GREY


def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text='', bold=False, italic=False, color=None, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(parts, style=None):
    """parts: lista di (text, dict) per run formattati misti."""
    p = doc.add_paragraph(style=style)
    for text, opts in parts:
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if opts.get('mono'):
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        if opts.get('color'):
            r.font.color.rgb = opts['color']
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.add_run(text)
    return p


def bullet_rich(parts, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    for text, opts in parts:
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if opts.get('mono'):
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
    return p


def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # bordo leggero
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '6')
        e.set(qn('w:color'), 'CCCCCC')
        pbdr.append(e)
    pPr.append(pbdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F6F8')
    pPr.append(shd)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], HEAD_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(cells[i], ALT_FILL)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def callout(label, text, fill="FFF4E5"):
    """Box evidenziato (es. Trappola / Nota / Best practice)."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(label + '  ')
    r.bold = True
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    doc.add_paragraph()  # spazio dopo


def spacer():
    doc.add_paragraph()


# =====================================================================
#  COPERTINA
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run('Scene Manager+')
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Manuale utente completo')
r.font.size = Pt(20)
r.font.color.rgb = GREY

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Plugin di gestione scene avanzata per SketchUp 2019')
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(60)
r = meta.add_run('Funzioni · Filosofia · Shortcut · Best practices · Limiti noti')
r.font.size = Pt(11)
r.font.color.rgb = DARK

doc.add_page_break()

# =====================================================================
#  INDICE (segnaposto - aggiornabile in Word con F9)
# =====================================================================
h1('Indice')
para('Per aggiornare i numeri di pagina: in Word, clic destro sull’indice → "Aggiorna campo" → "Aggiorna intero sommario". '
     'Se l’indice appare vuoto, premere Ctrl+A e poi F9.', italic=True, color=GREY, size=10)

# Inserisce un vero campo TOC
p = doc.add_paragraph()
run = p.add_run()
fldStart = OxmlElement('w:fldChar'); fldStart.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
fldText = OxmlElement('w:t'); fldText.text = "Aggiorna questo campo per generare l’indice."
fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
for el in (fldStart, instr, fldSep, fldText, fldEnd):
    run._r.append(el)

doc.add_page_break()

# =====================================================================
#  1. INTRODUZIONE E FILOSOFIA
# =====================================================================
h1('1. Introduzione e filosofia')

para('Scene Manager+ è un plugin per SketchUp 2019 che porta la gestione delle scene a un livello '
     'paragonabile ai "livelli di Photoshop": una lista riordinabile, organizzabile in cartelle, con '
     'esportazione batch delle immagini, watermark logo, cartiglio professionale e gestione avanzata degli stili.')

h2('1.1 Perché esiste')
para('SketchUp 2019 offre solo una barra di tab scene piatta, senza riordino reale, senza cartelle e senza '
     'un flusso di export strutturato. Scene Manager+ colma queste lacune restando dentro i limiti dell’API '
     'Ruby di SketchUp 2019 (molte funzioni native non sono esposte allo scripting).')

h2('1.2 Principi guida (filosofia)')
bullet_rich([('"Scatta foto completa". ', {'bold': True}),
             ('Quando crei o aggiorni una scena, il plugin cattura sempre lo stato completo del viewport '
              '(camera, stile, fog, layer, ecc.), non i "default scene properties" personalizzati di SketchUp. '
              'Il flusso è fotografico, non "rispetta i miei default".', {})])
bullet_rich([('Ordine logico, non distruttivo. ', {'bold': True}),
             ('L’ordine delle scene vive come metadato nel file .skp, senza mai cancellare e ricreare le '
              'pagine native (operazione rischiosa in SU 2019). Le tab native restano nell’ordine di creazione: '
              'è normale e voluto.', {})])
bullet_rich([('Un solo Ctrl+Z per azione. ', {'bold': True}),
             ('Ogni operazione del plugin è racchiusa in una singola operazione SketchUp: un annullamento '
              'pulito ripristina lo stato precedente.', {})])
bullet_rich([('Rispetto delle performance. ', {'bold': True}),
             ('Su file con observer di plugin terzi ogni scrittura può costare secondi. Il plugin minimizza '
              'le scritture (defer mode, cache in RAM, diff prima di scrivere).', {})])
bullet_rich([('Sviluppo a fasi con verifica. ', {'bold': True}),
             ('Le feature sono state aggiunte in fasi successive, ciascuna verificata sul campo.', {})])

h2('1.3 Tutto il dato viaggia col file')
para('Ordine, cartelle, colori, nickname degli stili, settings di export e naming: tutto è salvato come '
     'attributo del modello .skp. Aprendo lo stesso file su un’altra postazione ritrovi la tua organizzazione. '
     'Solo pochissime preferenze di macchina usano i default globali di SketchUp.')

# =====================================================================
#  2. INSTALLAZIONE E AVVIO
# =====================================================================
h1('2. Installazione e avvio')

h2('2.1 Avvio del plugin')
bullet_rich([('Menu ', {}), ('Plugins → Scene Manager+', {'bold': True}), (' oppure l’icona nella toolbar.', {})])
para('La finestra principale è una palette (stile "Utility"): resta sempre sopra il viewport e ricorda '
     'posizione e dimensione tra le sessioni. Se era aperta alla chiusura di SketchUp, viene riaperta '
     'automaticamente all’avvio successivo.')

h2('2.2 Aggiornamento del codice')
para('Dopo modifiche al codice del plugin i file vengono copiati nella cartella Plugins di SketchUp. Regola pratica:')
bullet_rich([('Modifiche Ruby (.rb): ', {'bold': True}), ('riavviare SketchUp.', {})])
bullet_rich([('Modifiche HTML/CSS/JS: ', {'bold': True}), ('basta chiudere e riaprire la finestra del plugin (c’è cache-busting).', {})])

# =====================================================================
#  3. INTERFACCIA PRINCIPALE
# =====================================================================
h1('3. L’interfaccia principale')

para('La finestra mostra la lista delle scene in ordine logico, eventualmente raggruppate in cartelle. '
     'Ogni riga rappresenta una scena con diversi elementi interattivi.')

h2('3.1 Anatomia di una riga scena')
table(
    ['Elemento', 'Funzione'],
    [
        ['Maniglia (grip)', 'Punto di trascinamento per riordinare. Un marker giallo indica la scena attiva nel viewport.'],
        ['Lettera stile (A, B, C…)', 'Badge che indica lo stile assegnato. Clic sx = Mini Style Manager; clic dx = riassegna stile.'],
        ['Checkbox export', 'Includi/escludi la scena dall’export (scope "all"). Persistente col file.'],
        ['Indice', 'Posizione 1-based nell’ordine logico globale (mantiene i "buchi" se escludi scene).'],
        ['Nome scena', 'Doppio clic = Properties dialog. Invio/F2 = rinomina inline.'],
        ['Swatch colore', 'Colore cosmetico per-scena. Clic sx = picker; clic dx = rimuovi colore.'],
        ['Thumbnail (opzionale)', 'Anteprima inline attivabile col toggle "Thumbs".'],
    ],
    widths=[1.8, 5.0]
)

h2('3.2 Selezione vs scena attiva')
para('Sono due concetti distinti, da non confondere:')
bullet_rich([('Selezione (barra azzurra): ', {'bold': True}),
             ('una o più scene evidenziate su cui agiscono i comandi (update, delete, export, ecc.).', {})])
bullet_rich([('Scena attiva (marker giallo): ', {'bold': True}),
             ('la singola scena effettivamente mostrata nel viewport.', {})])
para('Cambiare la selezione non cambia necessariamente la scena attiva, e viceversa. In defer mode il marker '
     'giallo non segue la selezione (il viewport non cambia).')

h2('3.3 Barra strumenti (toolbar)')
table(
    ['Pulsante', 'Azione'],
    [
        ['\U0001F4F7+ Nuova scena', 'Crea una scena dalla vista corrente, catturando lo stato completo del viewport.'],
        ['⟳ Update', 'Aggiorna la/le scena/e selezionata/e dalla vista corrente. Clic dx = "Update only…" (parziale).'],
        ['Defer', 'Attiva/disattiva la modalità differita (scritture in RAM, un solo flush).'],
        ['Thumbs', 'Mostra/nasconde le anteprime inline nelle righe.'],
        ['Cartella +', 'Crea una nuova cartella logica.'],
        ['Export', 'Apre il dialog di esportazione batch.'],
        ['Settings', 'Apre le impostazioni (naming, export, logo, interfaccia, stili).'],
    ],
    widths=[2.2, 4.6]
)

# =====================================================================
#  4. GESTIONE SCENE
# =====================================================================
h1('4. Gestione delle scene')

h2('4.1 Creare una scena dalla vista')
para('L’icona \U0001F4F7+ crea una nuova scena dalla vista corrente. A differenza del comando nativo, il plugin '
     'forza la cattura di tutti gli 8 flag (camera, stile, ombre, geometria nascosta, layer, piani di sezione, '
     'assi, opzioni di rendering), così la scena nasce sempre "completa". Replica anche correttamente la '
     'visibilità dei layer effettivamente mostrata nel viewport (incluso il caso "Add visible tag" del plugin '
     'Layers Manager).')

h2('4.2 Aggiornare una scena (Update from view)')
para('Il pulsante ⟳ nella toolbar aggiorna la scena selezionata (anche selezione multipla) dalla vista '
     'corrente, rispettando i flag attivi della scena.')
callout('⚠ Stile "dirty"',
        'Se lo stile attivo ha modifiche non salvate, compare un dialog a tre pulsanti (Sì / No / Annulla) '
        'equivalente al "Warning - Scenes and Styles" nativo: Sì aggiorna lo stile e la scena; No permette di '
        'salvare come nuovo stile; Annulla salva tutto tranne lo stile. Senza questo passaggio le modifiche allo '
        'stile andrebbero perse silenziosamente.')

h3('Update parziale (clic destro su ⟳)')
para('Clic destro sul pulsante Update apre il menu "Update only…" con le sole proprietà attive della scena. '
     'Selezionando una voce aggiorni solo quella proprietà (es. solo la camera, solo i layer), lasciando '
     'intatto il resto.')

h2('4.3 Selezione e navigazione da tastiera')
table(
    ['Tasto', 'Azione'],
    [
        ['PaginaSu / PaginaGiù', 'Sposta la selezione su/giù lungo l’ordine logico visibile (le cartelle chiuse vengono saltate).'],
        ['Home / Fine', 'Salta alla prima / ultima scena visibile.'],
        ['Freccia Su / Freccia Giù', 'SPOSTA la scena selezionata (riordino) nell’ordine logico.'],
        ['Invio / F2', 'Avvia la rinomina inline della scena selezionata.'],
        ['Esc', 'Annulla la rinomina inline in corso.'],
    ],
    widths=[2.3, 4.7]
)
callout('Nota sugli shortcut',
        'Queste scorciatoie funzionano solo dopo un clic dentro la finestra del plugin. SketchUp 2019 instrada '
        'quasi sempre i tasti al viewport, quindi non è possibile intercettarli globalmente dal dialog. Per uno '
        'shortcut "vero" col viewport in focus, vedi il comando "Jump to active scene" (cap. 13).')

h2('4.4 Riordino (drag & drop)')
para('Trascina la maniglia di una scena per spostarla. Puoi trascinare scene dentro/fuori dalle cartelle. '
     'Il drag & drop è implementato in modo custom (non HTML5 nativo) per funzionare in modo affidabile in CEF.')
callout('⏱ Primo riordino su file "vergine"',
        'La prima volta che riordini o crei una cartella su un file mai toccato dal plugin, c’è un breve freeze '
        '(qualche secondo su modelli pesanti): il plugin assegna gli identificativi univoci (uid) alle pagine. '
        'Le operazioni successive sono immediate.', fill="E8F4EA")

h2('4.5 Rinominare e descrivere')
bullet('Rinomina inline: seleziona la scena e premi Invio/F2, oppure clic destro → Rename. Invio conferma, Esc annulla.')
bullet('Properties dialog (doppio clic): permette di modificare nome e descrizione con commit immediato (senza pulsante Apply).')

h2('4.6 Eliminare scene')
para('Seleziona una o più scene ed elimina. In defer mode la cancellazione è differita fino al flush.')

# =====================================================================
#  5. CARTELLE
# =====================================================================
h1('5. Cartelle')

para('Le cartelle sono raggruppamenti logici, salvati nel file. Non corrispondono a pagine native di SketchUp: '
     'servono solo a organizzare la lista nel plugin.')
bullet('Crea una cartella dalla toolbar, assegnale un nome e un colore.')
bullet('Trascina le scene dentro/fuori dalle cartelle.')
bullet('Le cartelle si possono espandere/comprimere; le scene di cartelle chiuse vengono saltate dalla navigazione.')
bullet('Le scene non assegnate a nessuna cartella vivono al livello radice.')
para('Tecnicamente il modello dati supporta cartelle annidate, ma l’interfaccia attuale lavora su un solo '
     'livello di cartelle.')

# =====================================================================
#  6. DEFER MODE
# =====================================================================
h1('6. Modalità differita (Defer mode)')

para('Defer mode è una modalità pensata per i file pesanti o con observer di plugin terzi, dove ogni '
     'singola scrittura sul modello può costare secondi.')
h2('6.1 Come funziona')
bullet('Attivala col pulsante Defer in toolbar.')
bullet('Mentre è attiva, tutte le modifiche (rinomine, descrizioni, flag, cancellazioni, riordino, cartelle) vivono in RAM.')
bullet('Un solo "flush" applica tutto al modello in un’unica operazione: un solo Ctrl+Z annulla l’intero blocco.')
bullet('Il flush avviene quando disattivi il defer o alla chiusura della finestra (auto-flush).')

h2('6.2 Eccezioni che restano immediate')
para('Alcune operazioni dipendono dallo stato "qui e ora" del viewport e quindi non possono essere differite:')
bullet('Selezione/attivazione scena (navigazione).')
bullet('Update from view (cattura il viewport corrente).')
bullet('Assegnazione stile (modifica lo stile selezionato e aggiorna la pagina).')
bullet('Modifiche live nel Mini Style Manager.')
callout('Quando usarla',
        'Attiva il defer quando devi fare molte modifiche di fila (riordino massiccio, rinomine, cartelle) su un '
        'file lento. Per singole operazioni occasionali non serve.', fill="E8F4EA")

# =====================================================================
#  7. STILI
# =====================================================================
h1('7. Gestione degli stili')

para('SketchUp 2019 non permette via API di creare, rinominare o clonare stili liberamente. Scene Manager+ '
     'aggira questi limiti con un sistema a tre livelli: badge lettera, Mini Style Manager, pool di slot + nickname.')

h2('7.1 Il badge lettera')
para('Ogni scena mostra una lettera (A, B, C…) che identifica lo stile assegnato. Le lettere sono assegnate in '
     'ordine alfabetico sul nome (o nickname) di tutti gli stili del modello.')
table(
    ['Interazione', 'Effetto'],
    [
        ['Clic sinistro', 'Apre il Mini Style Manager per quello stile.'],
        ['Clic destro', 'Apre il picker per riassegnare lo stile alla scena (lista A, B, C… con lo stile corrente in giallo).'],
        ['Badge "?"', 'Lo stile della scena non è tra quelli enumerati: tipico delle scene Match Photo (vedi cap. 12).'],
    ],
    widths=[2.0, 4.8]
)

h2('7.2 Mini Style Manager (clic sinistro sulla lettera)')
para('Un editor compatto dello stile, con feedback live nel viewport. Le modifiche si applicano e vengono '
     'committate nello stile persistente immediatamente. Gruppi disponibili:')
bullet_rich([('Edges (linee): ', {'bold': True}), ('modalità colore linee, ordinamento trasparenze, pulsante "→ 1" che imposta profili visibili a spessore 1.', {})])
bullet_rich([('Background: ', {'bold': True}), ('colore di sfondo, orizzonte on/off, colore cielo.', {})])
bullet_rich([('Display: ', {'bold': True}), ('geometria nascosta, assi del modello (pulsante toggle), piani e tagli di sezione.', {})])
bullet_rich([('Nickname e badge color: ', {'bold': True}), ('nell’header, per dare allo stile un nome amichevole e un colore al badge.', {})])
bullet_rich([('Apri Window → Styles nativo: ', {'bold': True}), ('pulsante "⧉" per il pannello stili nativo di SketchUp.', {})])

callout('Scope sempre "tutte le scene con questo stile"',
        'Modificare uno stile nel Mini Style Manager influenza tutte le scene che lo usano. Per un override su una '
        'singola scena: duplica lo stile in Window → Styles nativo, poi riassegnalo con clic destro sulla lettera. '
        'L’API di SU 2019 non permette di clonare uno stile via codice.')

h2('7.3 Riassegnare uno stile (clic destro sulla lettera)')
para('Apre un picker con tutti gli stili. In cima c’è la voce "+ New style…" per creare uno stile nuovo che '
     'cattura ciò che il viewport sta mostrando (vedi sotto). Selezionando uno stile esistente, la scena viene '
     'riassegnata immediatamente.')

h2('7.4 Pool di slot + nickname (creare stili nuovi)')
para('Poiché SU 2019 non crea stili da zero via API, il plugin usa 25 file di stile pre-confezionati ("slot"). '
     'Quando crei un nuovo stile con "+ New style…":')
numbered('Il plugin alloca il primo slot libero.')
numbered('Cattura le opzioni di rendering correnti del viewport (incluse modifiche non salvate).')
numbered('Ti chiede un nickname (nome amichevole, opzionale ma consigliato).')
numbered('Assegna il nuovo stile alla scena.')
para('Il nickname è un nome "di facciata" che vedi solo nel plugin: il pannello Styles nativo continua a mostrare '
     'il nome reale (es. "SM+ Slot 03"). I nickname viaggiano col file e devono essere univoci.')
callout('Pool esaurito',
        'Se finisci i 25 slot, il plugin avvisa. Per estenderli servono nuovi file slot rigenerati (operazione da '
        'sviluppatore). In pratica 25 stili nuovi per modello sono ampiamente sufficienti.')

h2('7.5 Badge color per-stile')
para('Puoi assegnare un colore allo sfondo del badge lettera, per riconoscere a colpo d’occhio gli stili. Il '
     'testo della lettera passa automaticamente a chiaro o scuro per restare leggibile sul colore scelto. Si '
     'imposta dal Mini Style Manager (riga "Badge color").')

h2('7.6 Pulizia stili non usati')
para('Settings → sezione "Style pool" → "Purge unused styles…" elenca e rimuove gli stili non referenziati da '
     'alcuna scena (con conferma; l’operazione non è annullabile). Pulisce anche nickname e colori orfani.')

# =====================================================================
#  8. COLORI PER-SCENA
# =====================================================================
h1('8. Colori per-scena')

para('Indipendente dal colore del badge stile, ogni scena può avere un proprio colore cosmetico, mostrato nello '
     'swatch a destra del nome. È puramente organizzativo (non cambia nulla nel viewport).')
bullet('Clic sinistro sullo swatch: apre il color picker HSV (con opzione "None" e colori recenti).')
bullet('Clic destro sullo swatch: rimuove il colore (scorciatoia).')
para('Il picker è un selettore HSV custom (quadrato saturazione/valore + slider tonalità + campo hex + 24 preset + '
     'ultimi 5 colori usati), costruito perché in CEF di SketchUp 2019 il controllo nativo <input type="color"> '
     'è inutilizzabile. Per il colore per-scena il picker applica il valore solo al rilascio, per non scatenare '
     'scritture continue su file lenti.')

# =====================================================================
#  9. PROPERTIES DIALOG E FOG
# =====================================================================
h1('9. Properties dialog (doppio clic) e Fog')

h2('9.1 Proprietà della scena')
para('Doppio clic su una scena apre il Properties dialog con commit live (nessun pulsante Apply): il nome si '
     'salva su Invio o uscendo dal campo, la descrizione uscendo dal campo, i flag al cambiamento. C’è anche un '
     'pulsante "Update from view" con clic destro per l’update parziale.')

h2('9.2 Controllo della Fog (nebbia)')
para('La sezione Fog replica Window → Fog nativo con controlli più precisi.')
callout('⚠ Le label "0%" e "100%" sono FISSE: indicano la densità, non la posizione',
        'Lo slider a due cursori rappresenta la DISTANZA dalla camera. Il cursore blu "0%" è dove la nebbia '
        'inizia (ancora trasparente); il cursore arancione "100%" è dove la nebbia diventa opaca. Tra i due la '
        'densità interpola. Non sono percentuali di posizione sull’asse.')
bullet_rich([('Unità: ', {'bold': True}), ('i valori sono convertiti nelle unità del modello (m, mm, cm, ft, in); '
              'lo slider si adatta alla scala del modello.', {})])
bullet_rich([('Vincolo scena attiva: ', {'bold': True}), ('la fog si può leggere/modificare solo sulla scena '
              'attualmente attiva nel viewport (è una proprietà del modello). Se apri Properties su una scena non '
              'attiva, i controlli sono disabilitati con un pulsante "Activate this scene".', {})])
bullet_rich([('Vincolo start < end: ', {'bold': True}), ('l’inizio nebbia non può superare la fine.', {})])
bullet_rich([('Input numerici: ', {'bold': True}), ('gli spinner nativi sono inutilizzabili in CEF SU 2019, quindi '
              'ci sono pulsanti − / + con auto-ripetizione (tieni premuto). Shift = passo ×10, Ctrl = passo ÷10.', {})])

# =====================================================================
#  10. NAMING
# =====================================================================
h1('10. Naming pattern (schema dei nomi)')

para('Settings → Naming definisce il pattern usato per i nomi dei file esportati e per la voce "Tavola nr." del '
     'cartiglio.')
code_block('{prefix}{sep}{nnn}{sep}{scene_name}')
bullet_rich([('prefix: ', {'bold': True}), ('può essere il nome del file .skp, un testo personalizzato (usato anche '
              'come "Cliente" nel cartiglio), oppure assente.', {})])
bullet_rich([('nnn: ', {'bold': True}), ('numerazione progressiva con zero-padding. Mantiene i "buchi": se escludi '
              'la scena 3 dall’export, i numeri restano 1, 2, 4, 5… (la posizione tavola resta stabile).', {})])
bullet_rich([('sep: ', {'bold': True}), ('il separatore tra le parti.', {})])
para('Il pulsante "Rename scenes now" applica il pattern come rinomina effettiva delle scene.')

# =====================================================================
#  11. EXPORT
# =====================================================================
h1('11. Esportazione batch')

para('Il cuore del plugin: esporta in blocco le scene come immagini PNG o JPG, con sovrapposizioni opzionali '
     '(logo watermark, etichetta filename, cartiglio).')

h2('11.1 Scope di esportazione')
table(
    ['Scope', 'Cosa esporta'],
    [
        ['All', 'Tutte le scene, filtrate dal checkbox "export incluso" di ciascuna riga.'],
        ['Selected', 'Solo le scene selezionate (ignora il checkbox export).'],
        ['Folders', 'Le scene delle cartelle scelte (ignora il checkbox export).'],
    ],
    widths=[1.8, 5.0]
)

h2('11.2 Cartella di output intelligente')
para('Se non specifichi una cartella di output, il plugin la sceglie accanto al file .skp:')
bullet_rich([('Caso A: ', {'bold': True}), ('non esiste "Immagini/" → la crea ed esporta lì.', {})])
bullet_rich([('Caso B: ', {'bold': True}), ('"Immagini/" esiste e contiene file → crea "Immagini/Superate/NN/", '
              'sposta lì i vecchi file, poi esporta in "Immagini" svuotata.', {})])
bullet_rich([('Caso C: ', {'bold': True}), ('"Immagini/" esiste vuota → esporta lì.', {})])
para('La cartella "Superate" vive dentro "Immagini/" per tenere pulita la radice del progetto. Se il file .skp '
     'non è salvato, il plugin chiede la cartella manualmente.')

h2('11.3 Logo watermark')
para('Un logo (PNG bundlato o personalizzato in Settings) viene composito sull’immagine, con ancoraggio, offset, '
     'scala e opacità configurabili.')

h2('11.4 Etichetta filename')
para('Stampa il nome del file (senza estensione) in basso a sinistra dell’immagine. Configurabile in Settings → '
     'gruppo "filename label": font, dimensione, grassetto, colore (hex + swatch + preset), offset, opacità. '
     'Il testo è renderizzato con font reali tramite un processo di sistema nascosto.')

h2('11.5 Cartiglio (title block)')
para('Aggiunge un cartiglio professionale SOTTO l’immagine (estende il canvas verso il basso, non copre il '
     'disegno). Layout a box:')
bullet('Cliente (dal prefix custom) / Oggetto (nome scena).')
bullet('Fase di progetto: Preliminare / Definitivo / Esecutivo.')
bullet('Tavola nr. (dalla numerazione del naming) / Data (oggi o override).')
bullet('Progetto / Disegnato e controllato da.')
bullet('Dati aziendali (da assets/titleblock/company.txt).')
bullet('Logo aziendale (da assets/titleblock/logo.jpg).')
para('I box si auto-dimensionano sul testo, i font si riducono automaticamente se la riga è troppo lunga, e le '
     'baseline sono allineate tra box anche con dimensioni diverse.')

h2('11.6 Moltiplicatore spessore linee')
para('Due moltiplicatori separati: uno per l’export (default 2.0) e uno per le anteprime (default 1.0), per '
     'rendere le linee più marcate nelle immagini finali.')

# =====================================================================
#  12. ANTEPRIME
# =====================================================================
h1('12. Anteprime (previews / thumbnails)')

para('Il plugin può generare anteprime PNG persistenti per scena, salvate per-modello sul disco dell’utente. '
     'Le anteprime inline si attivano col toggle "Thumbs". La generazione è asincrona con progress bar (una scena '
     'per tick, per non bloccare l’interfaccia).')
callout('Niente preview nel pannello di default',
        'Su esplicita richiesta, non c’è un riquadro anteprima grande che rallenterebbe il refresh. Le thumbnail '
        'inline sono un’opzione attivabile.')

# =====================================================================
#  13. SHORTCUT GLOBALE
# =====================================================================
h1('13. Comando "Jump to active scene" e shortcut globali')

para('In SketchUp 2019 una finestra HtmlDialog non riceve in modo affidabile i tasti quando il focus è sul '
     'viewport. Per uno shortcut "vero" (premere un tasto col viewport in primo piano), il plugin registra un '
     'comando di menu: Plugins → "Scene Manager+: Jump to active scene".')
para('Questo comando ri-applica camera/stile/layer della scena attiva (equivale a cliccare il nome della scena). '
     'Puoi assegnargli uno shortcut da Window → Preferences → Shortcuts (SU mostra lì solo i comandi che hanno '
     'una voce di menu).')

# =====================================================================
#  14. RIEPILOGO SHORTCUT
# =====================================================================
h1('14. Riepilogo scorciatoie')

para('Funzionano dopo un clic dentro la finestra del plugin.')
table(
    ['Scorciatoia', 'Azione'],
    [
        ['PaginaSu / PaginaGiù', 'Muove la selezione lungo la lista (salta cartelle chiuse).'],
        ['Home / Fine', 'Prima / ultima scena.'],
        ['Freccia Su / Freccia Giù', 'Sposta (riordina) la scena selezionata.'],
        ['Invio / F2', 'Rinomina inline la scena selezionata.'],
        ['Esc', 'Annulla rinomina inline.'],
        ['Doppio clic', 'Apre il Properties dialog.'],
        ['Clic sx lettera', 'Mini Style Manager.'],
        ['Clic dx lettera', 'Picker riassegna stile.'],
        ['Clic sx swatch colore', 'Color picker per-scena.'],
        ['Clic dx swatch colore', 'Rimuove colore per-scena.'],
        ['Clic dx su Update (⟳)', 'Menu "Update only…" (parziale).'],
        ['Ctrl+Z', 'Annulla l’ultima operazione del plugin (una per azione).'],
    ],
    widths=[2.3, 4.7]
)

# =====================================================================
#  15. BEST PRACTICES
# =====================================================================
h1('15. Best practices')

h2('15.1 Workflow consigliato')
numbered('Organizza le scene in cartelle per blocchi logici (sezioni, prospetti, dettagli…).')
numbered('Assegna nickname e badge color agli stili ricorrenti per riconoscerli a colpo d’occhio.')
numbered('Imposta il naming pattern e il cartiglio una volta sola: viaggiano col file.')
numbered('Prima di un export importante, verifica i checkbox "export incluso" delle scene.')
numbered('Usa lo scope "All" per l’export completo, "Selected"/"Folders" per export parziali.')

h2('15.2 Su file pesanti o con plugin terzi')
bullet('Attiva il Defer mode prima di una sessione di molte modifiche (riordino, rinomine, cartelle).')
bullet('Aspettati un freeze una tantum al primo riordino/cartella su un file mai toccato dal plugin (assegnazione uid).')
bullet('Evita di muovere di continuo i color picker se il colore non ha effetto live: applica al rilascio.')

h2('15.3 Stili')
bullet('Per modificare lo stile di una sola scena senza toccare le altre: duplica lo stile in Window → Styles nativo, poi riassegnalo via clic destro sulla lettera.')
bullet('Quando lo stile è "dirty", scegli con cura nel dialog Sì/No/Annulla: "No → Save as new" preserva lo stile originale.')
bullet('Periodicamente, fai "Purge unused styles" per ripulire gli slot non usati.')

h2('15.4 Backup')
bullet('Tutta l’organizzazione vive nel .skp: un backup del file preserva ordine, cartelle, colori, nickname e settings.')
bullet('Le anteprime su disco sono rigenerabili: non servono per il backup.')

# =====================================================================
#  16. LIMITI NOTI
# =====================================================================
h1('16. Limiti noti e casi particolari')

h2('16.1 Match Photo')
para('Le scene Match Photo (foto di sfondo) sono un caso speciale: l’API Ruby di SketchUp 2019 non espone quasi '
     'nulla del sottosistema Match Photo, quindi il plugin NON può clonarle preservando la foto.')
bullet('Creare una nuova scena partendo da una scena Match Photo non copia la foto: il plugin avvisa e suggerisce di usare il "+" dell’inspector nativo (Window → Scenes).')
bullet('Riassegnare lo stile a una scena Match Photo è bloccato (rischio crash + perdita della foto).')
bullet('Il badge lettera "?" su una scena è spesso il segnale che si tratta di una scena Match Photo.')
para('Per duplicare scene Match Photo, usare il pulsante "+" dell’inspector Scenes nativo di SketchUp.')

h2('16.2 Ordine delle tab native')
para('Le tab scene native di SketchUp restano nell’ordine di creazione: il riordino vive solo nell’ordine '
     'logico del plugin. È una scelta di progetto (riordinare le pagine native sarebbe distruttivo in SU 2019).')

h2('16.3 Shortcut globali')
para('Non è possibile intercettare i tasti globalmente dalla finestra del plugin (limite CEF di SU 2019). Usa il '
     'comando di menu "Jump to active scene" con shortcut assegnabile.')

h2('16.4 Moltiplicatore linee sulle anteprime')
para('Il moltiplicatore spessore linee funziona sull’export ma sembra inefficace sulle anteprime 300×150 '
     '(probabile clamping della larghezza a 1px alla bassa risoluzione). Limite noto, non bloccante.')

h2('16.5 Geometria nascosta e piani di sezione')
para('Stati legati alla geometria del file (hidden geometry, section planes) non sono trasferibili tra file e '
     'non sono inclusi nel (futuro) copia/incolla scene cross-file.')

# =====================================================================
#  17. RISOLUZIONE PROBLEMI
# =====================================================================
h1('17. Risoluzione problemi')

table(
    ['Sintomo', 'Causa / Soluzione'],
    [
        ['Le tab scene restano visibili dopo aver riaperto un file',
         'Bug noto SU 2019. Workaround: View → Scene Tabs, clicca due volte (accende e spegne davvero). Oppure '
         'attiva in Settings → Interface "Force Scene Tabs OFF".'],
        ['Una scena mostra il badge "?"',
         'Lo stile non è enumerato: tipicamente una scena Match Photo. Normale.'],
        ['Il marker giallo non segue la selezione',
         'Comportamento corretto in Defer mode (il viewport non cambia). In modalità normale il giallo segue la scena attiva.'],
        ['Un nuovo stile cattura le impostazioni sbagliate',
         '"+ New style…" fotografa il viewport corrente: assicurati che il viewport mostri ciò che vuoi prima di crearlo.'],
        ['Freeze di qualche secondo al primo riordino',
         'Assegnazione uid una tantum su file mai toccato dal plugin. Le operazioni successive sono immediate.'],
        ['Le modifiche allo stile spariscono',
         'Probabile cambio stile con modifiche "dirty" non salvate. Usa il dialog Sì/No/Annulla e l’opzione "Save as new".'],
        ['Il color picker nativo non si apre',
         'In CEF SU 2019 <input type=color> è inutilizzabile: il plugin usa picker HSV custom (già attivi ovunque).'],
    ],
    widths=[2.6, 4.4]
)

# =====================================================================
#  CHIUSURA
# =====================================================================
spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Fine del manuale —')
r.italic = True
r.font.color.rgb = GREY

out = r'C:\Claude\Sketchup Plugins\SceneManagerPlus\Scene Manager+ - Manuale utente.docx'
doc.save(out)
print('Salvato:', out)
